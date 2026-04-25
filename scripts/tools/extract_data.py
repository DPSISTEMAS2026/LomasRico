import pandas as pd
from docx import Document
from pypdf import PdfReader
import os

def read_excel(path):
    print(f"\n--- ANALYZING EXCEL: {path} ---")
    try:
        xls = pd.ExcelFile(path)
        print(f"Sheets: {xls.sheet_names}")
        for sheet in xls.sheet_names:
            print(f"\n[SHEET: {sheet}]")
            df = pd.read_excel(xls, sheet_name=sheet)
            # Print columns and first 20 rows formatted
            print(df.head(20).to_string(index=False))
    except Exception as e:
        print(f"Error reading Excel: {e}")

def read_docx(path):
    print(f"\n--- ANALYZING DOCX: {path} ---")
    try:
        doc = Document(path)
        print("TEXT CONTENT:")
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text.strip())
        
        print("\nTABLES CONTENT:")
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                print(" | ".join(cells))
            print("-" * 20)
    except Exception as e:
        print(f"Error reading DOCX: {e}")

def read_pdf(path):
    print(f"\n--- ANALYZING PDF: {path} ---")
    try:
        reader = PdfReader(path)
        for i, page in enumerate(reader.pages):
            print(f"\n[PAGE {i+1}]")
            print(page.extract_text())
    except Exception as e:
        print(f"Error reading PDF: {e}")

def main():
    base_path = "c:/Users/ddiaz/OneDrive/Escritorio/PRODUCCION LO MAS RICO V3/"
    files = {
        "excel": "MENU RESTAURANTE.xlsx",
        "docx": "Gramajesceciches.docx",
        "pdf": "RECETAS   LOMASRICO.pdf"
    }

    for ftype, fname in files.items():
        full_path = os.path.join(base_path, fname)
        if os.path.exists(full_path):
            if ftype == "excel": read_excel(full_path)
            elif ftype == "docx": read_docx(full_path)
            elif ftype == "pdf": read_pdf(full_path)
        else:
            print(f"File not found: {fname}")

if __name__ == "__main__":
    main()
