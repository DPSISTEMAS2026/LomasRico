import pandas as pd
from docx import Document
from pypdf import PdfReader
import os
import sys

def read_excel(path, f):
    f.write(f"\n\n{'='*20}\nANALYZING EXCEL: {path}\n{'='*20}\n")
    try:
        xls = pd.ExcelFile(path)
        f.write(f"Sheets: {xls.sheet_names}\n")
        for sheet in xls.sheet_names:
            f.write(f"\n[SHEET: {sheet}]\n")
            df = pd.read_excel(xls, sheet_name=sheet)
            f.write(df.to_string(index=False))
            f.write("\n")
    except Exception as e:
        f.write(f"Error reading Excel: {e}\n")

def read_docx(path, f):
    f.write(f"\n\n{'='*20}\nANALYZING DOCX: {path}\n{'='*20}\n")
    try:
        doc = Document(path)
        f.write("TEXT CONTENT:\n")
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text.strip() + "\n")
        
        f.write("\nTABLES CONTENT:\n")
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                f.write(" | ".join(cells) + "\n")
            f.write("-" * 20 + "\n")
    except Exception as e:
        f.write(f"Error reading DOCX: {e}\n")

def read_pdf(path, f):
    f.write(f"\n\n{'='*20}\nANALYZING PDF: {path}\n{'='*20}\n")
    try:
        reader = PdfReader(path)
        for i, page in enumerate(reader.pages):
            f.write(f"\n[PAGE {i+1}]\n")
            try:
                text = page.extract_text()
                f.write(text if text else "[No text extracted]")
            except Exception as e:
                f.write(f"Error extracting text from page {i}: {e}")
            f.write("\n")
    except Exception as e:
        f.write(f"Error reading PDF: {e}\n")

def main():
    base_path = "c:/Users/ddiaz/OneDrive/Escritorio/PRODUCCION LO MAS RICO V3/"
    files = {
        "excel": "MENU RESTAURANTE.xlsx",
        "docx": "Gramajesceciches.docx",
        "pdf": "RECETAS   LOMASRICO.pdf"
    }

    with open("analysis_output.txt", "w", encoding="utf-8") as f:
        for ftype, fname in files.items():
            full_path = os.path.join(base_path, fname)
            if os.path.exists(full_path):
                if ftype == "excel": read_excel(full_path, f)
                elif ftype == "docx": read_docx(full_path, f)
                elif ftype == "pdf": read_pdf(full_path, f)
            else:
                f.write(f"File not found: {fname}\n")

if __name__ == "__main__":
    main()
